import re
import streamlit as st
from docx import Document
from datetime import datetime
import openai
import os
from dotenv import load_dotenv

# Load environment variables and set OpenAI API key
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Function to read file content with appropriate encoding
def read_file(uploaded_file):
    if uploaded_file.name.endswith(".txt"):
        try:
            # Try reading as UTF-8
            content = uploaded_file.read().decode("utf-8")
        except UnicodeDecodeError:
            # Fallback to ISO-8859-1 if UTF-8 fails
            uploaded_file.seek(0)
            content = uploaded_file.read().decode("ISO-8859-1")
    elif uploaded_file.name.endswith(".docx"):
        # Use python-docx for .docx files
        doc = Document(uploaded_file)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        st.error("Unsupported file type. Please upload a .txt or .docx file.")
        content = None
    return content

# Function to chunk text into manageable parts
def chunk_text(text, max_tokens=80000):
    words = text.split()
    chunks = []
    current_chunk = []
    current_tokens = 0

    for word in words:
        current_tokens += len(word) // 4  # Approximate token count
        current_chunk.append(word)
        
        if current_tokens >= max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_tokens = 0

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

# Function to generate responses from OpenAI with adjustable parameters
def openai_prompt(prompt, temperature, top_p, frequency_penalty, presence_penalty):
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are the best business coach summary transcriber."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        temperature=temperature,
        top_p=top_p,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty
    )
    summary_text = response.choices[0].message.content.strip()
    return summary_text

# Process each chunk for specific prompts
def process_chunks(transcript, prompt_text, temperature, top_p, frequency_penalty, presence_penalty):
    chunks = chunk_text(transcript)
    results = []
    for chunk in chunks:
        prompt = f"{prompt_text}\n\n{chunk}"
        result = openai_prompt(prompt, temperature, top_p, frequency_penalty, presence_penalty)
        results.append(result)
    return " ".join(results)

# Function to clean Markdown-style formatting
def clean_text(text):
    # Remove Markdown symbols like ##, ###, **, and -
    clean_text = re.sub(r"[#\*\-]+", "", text)
    # Replace multiple newlines with a single newline
    clean_text = re.sub(r"\n\s*\n", "\n", clean_text)
    return clean_text.strip()

# Function to analyze transcript
def generate_summary(transcript, action_points_prompt, recommendations_prompt, temperature, top_p, frequency_penalty, presence_penalty):
    # 1) Main action points
    main_points = process_chunks(transcript, action_points_prompt, temperature, top_p, frequency_penalty, presence_penalty).splitlines()

    # 2) Recommendations
    recommendations = process_chunks(transcript, recommendations_prompt, temperature, top_p, frequency_penalty, presence_penalty)

    # Clean up formatting in main_points and recommendations
    main_points = [clean_text(point) for point in main_points]
    recommendations = clean_text(recommendations)

    return main_points, recommendations

# Function to export as Word document
def export_to_word(main_points, recommendations):
    doc = Document()

    doc.add_heading("Main Discussion Points", level=1)
    for point in main_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph(recommendations)

    doc_name = f"Meeting_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(doc_name)
    return doc_name

# Streamlit App Layout
st.title("Meeting Notes Summarizer")

uploaded_file = st.file_uploader("Upload a .txt or .docx file with meeting notes", type=["txt", "docx"])

# Sidebar inputs for prompts and OpenAI parameters
st.sidebar.header("Custom Prompts")
action_points_prompt = st.sidebar.text_area(
    "Action Points Prompt",
    """Extract Basic Information
Title
Date/Time
Participants
Generate a Concise Meeting Summary
Summarize the main discussion points in 150 words or fewer.
Include strategic priorities, key decisions, and important challenges."""
)
recommendations_prompt = st.sidebar.text_area(
    "Recommendations Prompt",
    """Identify Business Insights and Challenges
Review the transcript for insights, opportunities and challenges facing the client.
 
Identify Outputs & Outcomes
Review the transcript for any mention of the following Outputs & Outcomes:
Accessed New Markets
Debt Finance Raised
EEN Network Advisory Achievement
Grant Funding Secured
Investment Raised
Jobs Created
Jobs Maintained
Partnering Achievement
Cross-reference each output and include it in the "Outputs & Outcomes" section if mentioned.
 
Identify Referrals to Other Support
Review the transcript for referrals and match against the following categories:
DBT (DIT)
Local Service (Growth Hub)
Knowledge Base/Universities
Knowledge Transfer Network
National Enquiry Gateway
Catapult
Private Sector Support
Scale Up Programme EOI
Peer-to-Peer Network
Other (please state)
 
Identify Key Actions
Extract no more than five clear and actionable steps for participants to progress.
 
Format:
Structure the output as follows:
 
Business Insights/Opportunities/Challenges: [Generate a bullet point list presenting each I/O/C in bold followed by a : (colon) and then an explanation]
 
+

Outputs & Outcomes: 
[Generate a bullet point list presenting each Output and Outcome in bold followed by a : (colon) and then an explanation]
 
+

Referrals to Other Support: 
[Generate a bullet point list presenting each I/O/C in bold followed by a : (colon) and then an explanation]
 
+

Actions:
[Generate a numbered bullet point list of no more than five actions presenting each Action/Activity in bold followed by a : (colon) and then an explanation]
"""
)

st.sidebar.header("OpenAI Settings")
temperature = st.sidebar.slider("Temperature", 0.0, 2.0, 0.5, step=0.1)
top_p = st.sidebar.slider("Top P", 0.0, 1.0, 1.0, step=0.1)
frequency_penalty = st.sidebar.slider("Frequency Penalty", 0.0, 2.0, 0.0, step=0.1)
presence_penalty = st.sidebar.slider("Presence Penalty", 0.0, 2.0, 0.0, step=0.1)

if uploaded_file:
    # Load file content
    transcript = read_file(uploaded_file)
    
    if transcript:
        # Run extraction and display results
        main_points, recommendations = generate_summary(
            transcript, action_points_prompt, recommendations_prompt, temperature, top_p, frequency_penalty, presence_penalty
        )

        with st.expander("Main Discussion Points"):
            st.write("\n".join(main_points))
        
        with st.expander("Recommendations"):
            st.write(recommendations)

        if st.button("Download Summary as Word Document"):
            doc_name = export_to_word(main_points, recommendations)
            with open(doc_name, "rb") as f:
                st.download_button("Download Word Document", data=f, file_name=doc_name)

st.write("Please upload your meeting notes to generate a summary.")
